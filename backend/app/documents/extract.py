"""Text extraction from documents (PDF, Word, images, plain text).

Pure-Python extractors with lazy, optional dependencies:

* PDF  → ``pypdf``      (install ``.[documents]``)
* DOCX → ``python-docx`` (install ``.[documents]``)
* images → no local OCR; flagged for the vision model in the pipeline
* text/* → decoded directly

Extraction never executes document content; it only reads text. Extracted text is treated as
UNTRUSTED external content downstream (content-trust), never as instructions.
"""

from __future__ import annotations

import mimetypes
from dataclasses import dataclass
from enum import Enum
from io import BytesIO

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".tiff", ".tif", ".bmp", ".webp", ".heic"}


class DocKind(str, Enum):
    pdf = "pdf"
    word = "word"
    image = "image"
    text = "text"
    unknown = "unknown"


@dataclass
class ExtractResult:
    kind: DocKind
    text: str
    needs_vision: bool = False  # image or scanned PDF: needs a multimodal model to understand


def detect_kind(filename: str, content_type: str | None = None) -> DocKind:
    name = (filename or "").lower()
    ctype = (content_type or mimetypes.guess_type(name)[0] or "").lower()
    if name.endswith(".pdf") or ctype == "application/pdf":
        return DocKind.pdf
    if name.endswith((".docx", ".doc")) or "word" in ctype or "officedocument.wordprocessing" in ctype:
        return DocKind.word
    if any(name.endswith(ext) for ext in _IMAGE_EXTS) or ctype.startswith("image/"):
        return DocKind.image
    if ctype.startswith("text/") or name.endswith((".txt", ".md", ".csv")):
        return DocKind.text
    return DocKind.unknown


def _extract_pdf(data: bytes) -> ExtractResult:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ExtractResult(DocKind.pdf, "", needs_vision=False)
    reader = PdfReader(BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    text = "\n".join(p for p in parts if p).strip()
    # Very little text from a PDF usually means a scanned/image PDF → needs vision.
    return ExtractResult(DocKind.pdf, text, needs_vision=len(text) < 20)


def _extract_docx(data: bytes) -> ExtractResult:
    try:
        import docx
    except ImportError:
        return ExtractResult(DocKind.word, "")
    document = docx.Document(BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    return ExtractResult(DocKind.word, "\n".join(parts).strip())


def _extract_text(data: bytes) -> ExtractResult:
    return ExtractResult(DocKind.text, data.decode("utf-8", errors="replace").strip())


def extract_text(data: bytes, filename: str, content_type: str | None = None) -> ExtractResult:
    """Extract plain text from a document's bytes based on its detected kind."""
    kind = detect_kind(filename, content_type)
    if kind is DocKind.pdf:
        return _extract_pdf(data)
    if kind is DocKind.word:
        return _extract_docx(data)
    if kind is DocKind.image:
        return ExtractResult(DocKind.image, "", needs_vision=True)
    if kind is DocKind.text:
        return _extract_text(data)
    # Unknown: best-effort decode, but don't claim confidence.
    try:
        return ExtractResult(DocKind.unknown, data.decode("utf-8", errors="strict").strip())
    except UnicodeDecodeError:
        return ExtractResult(DocKind.unknown, "", needs_vision=False)
